from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def update_thesis():
    doc = Document('thesis.docx')
    
    optimized_content = []
    with open('optimized_thesis_content.txt', 'r', encoding='utf-8') as f:
        for line in f:
            optimized_content.append(line.rstrip('\n'))
    
    for p in doc.paragraphs:
        p.text = ''
    
    for i in range(len(doc.paragraphs)):
        if i < len(optimized_content):
            doc.paragraphs[i].text = optimized_content[i]
        else:
            if i < len(doc.paragraphs):
                doc.paragraphs[i].text = ''
    
    doc.save('thesis_optimized.docx')
    print("Optimized thesis saved as thesis_optimized.docx")

if __name__ == '__main__':
    update_thesis()
